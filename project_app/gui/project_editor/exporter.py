from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

def export_project_to_word(file_path, projekt_data):
    """
    Egy projekt teljes adatstruktúráját formázott Word dokumentumba menti.
    :param file_path: Hová mentse a .docx fájlt.
    :param projekt_data: Dict a projekt minden adatával.
    """
    doc = Document()

    # --- Címlap, fejléc ---
    doc.add_heading("VILLAMOS KIVITELEZÉSI ÁRAJÁNLAT/JEGYZŐKÖNYV", level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Dátum: {datetime.now().strftime('%Y.%m.%d.')}")
    doc.add_paragraph(f"Projekt neve: {projekt_data['projektnev']}")
    doc.add_paragraph(f"Megrendelő neve: {projekt_data['megrendelo']}")
    doc.add_paragraph(f"Lakcím: {projekt_data['cim']}")
    doc.add_paragraph("")

    osszeg_total = 0  # Összesítés

    # --- Helyiségek feldolgozása ---
    for helyseg in projekt_data.get('helyisegek', []):
        doc.add_heading(f"Helyiség: {helyseg.get('nev','')}", level=1)

        # --- Anyag tételek táblázat ---
        anyagok = helyseg.get('tetel_lista', [])
        if anyagok:
            doc.add_paragraph("Anyag tételek:", style='Intense Quote')
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Tétel"
            hdr_cells[1].text = "Egység"
            hdr_cells[2].text = "Mennyiség"
            hdr_cells[3].text = "Egységár (Ft)"
            hdr_cells[4].text = "Összesen (Ft)"
            hdr_cells[5].text = "Megjegyzés"
            helyseg_osszeg = 0

            for t in anyagok:
                row = table.add_row().cells
                row[0].text = t.get('nev', '')
                row[1].text = t.get('egyseg', '')
                row[2].text = str(t.get('mennyiseg', ''))
                row[3].text = str(t.get('egysegar', ''))
                try:
                    ossz = int(t.get('mennyiseg',1)) * float(t.get('egysegar',0))
                except Exception:
                    ossz = 0
                row[4].text = f"{ossz:,.0f}"
                row[5].text = t.get('megjegyzes', '')
                helyseg_osszeg += ossz

            doc.add_paragraph("")  # Szóköz

        # --- Kivitelezési tételek táblázat ---
        kivitek = helyseg.get('kivitelezesi_tetelek', [])
        if kivitek:
            doc.add_paragraph("Kivitelezési munkák:", style='Intense Quote')
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Munka"
            hdr_cells[1].text = "Egység"
            hdr_cells[2].text = "Mennyiség"
            hdr_cells[3].text = "Egységár (Ft)"
            hdr_cells[4].text = "Összesen (Ft)"
            hdr_cells[5].text = "Megjegyzés"

            for t in kivitek:
                row = table.add_row().cells
                row[0].text = t.get('nev', '')
                row[1].text = t.get('egyseg', '')
                row[2].text = str(t.get('mennyiseg', ''))
                row[3].text = str(t.get('egysegar', ''))
                try:
                    ossz = int(t.get('mennyiseg',1)) * float(t.get('egysegar',0))
                except Exception:
                    ossz = 0
                row[4].text = f"{ossz:,.0f}"
                row[5].text = t.get('megjegyzes', '')
                helyseg_osszeg += ossz

        # --- Részösszeg helyiségenként ---
        if anyagok or kivitek:
            doc.add_paragraph(f"Helyiség részösszeg: {helyseg_osszeg:,.0f} Ft", style='Normal')
            doc.add_paragraph("")
            osszeg_total += helyseg_osszeg

    # --- Végösszeg, aláírás rész ---
    doc.add_paragraph("")
    doc.add_heading(f"Projekt végösszeg: {osszeg_total:,.0f} Ft", level=2)
    doc.add_paragraph("")
    doc.add_paragraph("_____________________________\nKivitelező/Átadó aláírása", style='Intense Quote')
    doc.add_paragraph("_____________________________\nMegrendelő/Átvevő aláírása", style='Intense Quote')

    doc.save(file_path)
